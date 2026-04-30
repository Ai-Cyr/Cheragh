"""Optional DOCX loader."""
from __future__ import annotations

from pathlib import Path
import hashlib

from ...base import Document
from .text import _file_metadata


def load_docx_file(path: str | Path) -> Document:
    """Load a DOCX document.

    Requires the optional dependency ``python-docx``:
    ``pip install cheragh[docx]``.
    """
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ImportError("DOCX loading requires python-docx. Install with: pip install cheragh[docx]") from exc

    p = Path(path)
    document = docx.Document(str(p))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    text = "\n\n".join(paragraphs)
    return Document(
        content=text,
        doc_id=_stable_file_doc_id(p),
        metadata=_file_metadata(p),
    )


def _stable_file_doc_id(path: Path) -> str:
    digest = hashlib.sha1(str(path.resolve()).encode("utf-8")).hexdigest()[:12]
    return f"file-{digest}"
